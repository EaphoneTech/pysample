import contextlib
import importlib.util
import itertools
import json
import logging
import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Any

import pathspec
import typer
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

logging.basicConfig(
    level=logging.INFO,
    format="%(levelname)-8s | %(message)s",
)
logger = logging.getLogger("pysample")

app = typer.Typer()

# 配置文件路径（与 bring-it/sample 兼容）
CONFIG_PATH = Path(".bring-it/sample.config.json")

# 自动探测用的「语言名 -> 扩展名」映射
LANGUAGE_MAP: dict[str, list[str]] = {
    "Python": ["py"],
    "JavaScript": ["js", "jsx", "mjs", "cjs"],
    "TypeScript": ["ts", "tsx"],
    "Java": ["java"],
    "Go": ["go"],
    "Rust": ["rs"],
    "C": ["c", "h"],
    "C++": ["cpp", "cc", "cxx", "hpp", "hxx"],
    "C#": ["cs"],
    "Ruby": ["rb"],
    "PHP": ["php"],
    "HTML": ["html", "htm"],
    "CSS": ["css"],
    "Shell": ["sh", "bash"],
    "Kotlin": ["kt", "kts"],
    "Swift": ["swift"],
    "Scala": ["scala"],
    "Lua": ["lua"],
    "SQL": ["sql"],
    "Vue": ["vue"],
}

# 源码探测时跳过的常见垃圾/产物目录
SKIP_DIRS = {
    ".git",
    "node_modules",
    "__pycache__",
    ".venv",
    "venv",
    "dist",
    "build",
    ".mypy_cache",
    ".pytest_cache",
    ".ruff_cache",
    "target",
    "out",
    ".idea",
    ".vscode",
}


def read_config(config_path: Path) -> dict[str, Any]:
    try:
        with open(config_path, encoding="utf-8") as file:
            config = json.load(file)
        logger.info("已从 %s 加载了配置文件", config_path)
        return config
    except Exception:
        logger.error("无法从 %s 加载配置文件", config_path.resolve())
        raise


def scan_files(
    base_dir: str,
    patterns: list[str],
    extensions: list[str],
    ignore_dirs: list[str],
    verbose: bool,
    title: str,
) -> list[Path]:
    base_path = Path(base_dir)
    if not base_path.exists() or not base_path.is_dir():
        logger.error("%s: 基础路径 %s 不存在或不是一个目录", title, base_path)
        raise ValueError(f"基础路径 {base_path} 不存在或不是一个目录")

    if patterns is None or len(patterns) == 0:
        logger.error("%s: 未指定扫描模式 'patterns'", title)
        raise ValueError("未指定扫描模式")

    if extensions is None or len(extensions) == 0:
        logger.error("%s: 未指定文件后缀 'extensions'", title)
        raise ValueError("未指定文件后缀")

    # 在这里将 patterns 拆分为 已经是文件的 , 和需要 glob 的
    matched_files: list[Path] = []
    glob_patterns: list[str] = []

    for p in patterns:
        if p is not None:
            pp = Path(p)
            if pp.exists() and pp.is_file():
                if verbose:
                    logger.debug("直接命中文件: %s with pattern: %s", base_path, p)
                matched_files.append(pp)
                continue
        glob_patterns.append(p)

    # 创建一个包含手动指定 ignore_dirs 的 pathspec 规则列表
    ignore_patterns = ignore_dirs[:]  # 复制 ignore_dirs 列表

    # 尝试读取 .gitignore 文件并添加到忽略规则中
    gitignore_path = base_path / ".gitignore"
    if gitignore_path.exists():
        with open(gitignore_path, encoding="utf-8") as f:
            gitignore_content = f.read()
        # 过滤掉空行和注释行
        gitignore_patterns = [
            line.strip()
            for line in gitignore_content.splitlines()
            if line.strip() and not line.startswith("#")
        ]
        ignore_patterns.extend(gitignore_patterns)

    # 使用 pathspec 创建忽略规则
    ignore_spec = pathspec.PathSpec.from_lines("gitwildmatch", ignore_patterns)

    for pattern, ext in itertools.product(glob_patterns, extensions):
        if not pattern.endswith("/"):
            pattern += "/"
        glob_pattern_str = f"{pattern}**/*.{ext}" if pattern else f"**/*.{ext}"
        if verbose:
            logger.debug(
                "正在扫描路径: %s with pattern: %s", base_path, glob_pattern_str
            )
        for file_path in base_path.glob(glob_pattern_str):
            if file_path.is_file():
                # 计算相对于 base_path 的路径用于匹配忽略规则
                relative_path = file_path.relative_to(base_path)

                # 检查是否匹配忽略规则
                if not ignore_spec.match_file(str(relative_path)):
                    matched_files.append(relative_path)
                    logger.debug("匹配的文件: %s", relative_path)

    logger.info("总匹配文件数: %d", len(matched_files))
    return matched_files


def _prepare_temp_directory(matched_files: list[Path], base_dir: Path) -> Path:
    """将扫描命中的文件复制到临时目录，保持相对目录结构。

    Returns:
        临时目录的 Path。
    """
    temp_dir = Path(tempfile.mkdtemp(prefix="pysample_"))
    base_resolved = base_dir.resolve()

    for f in matched_files:
        # 解析为绝对路径
        abs_f = f if f.is_absolute() else (base_resolved / f).resolve()

        # 尝试获取相对于 base_dir 的相对路径
        try:
            rel = abs_f.relative_to(base_resolved)
        except ValueError:
            # 文件在 base_dir 外部时，回退到仅用文件名
            rel = Path(abs_f.name)

        dest = temp_dir / rel
        dest.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(abs_f, dest)
        logger.debug("暂存文件: %s -> %s", abs_f, dest)

    return temp_dir


def _write_extractor_config(
    temp_dir: Path, output_file: Path, group: dict[str, Any]
) -> Path:
    """生成 copyright-code-extractor 所需的 .copyright-extractor.json 文件。

    Returns:
        配置文件的 Path。
    """
    max_lines = group.get("maxLines", 0)
    config: dict[str, Any] = {
        "project_root": str(temp_dir),
        "output_file": str(output_file.resolve()),
        "software_name": group.get("title", "未命名文档"),
        "software_version": group.get("version", "V1.0"),
    }

    if max_lines > 0:
        config["lines_to_extract"] = max_lines
    else:
        config["extract_all"] = True

    config_path = temp_dir / ".copyright-extractor.json"
    with open(config_path, "w", encoding="utf-8") as f:
        json.dump(config, f, indent=2, ensure_ascii=False)
    logger.debug("已生成配置文件: %s", config_path)
    return config_path


def remove_consecutive_blank_lines(text: str) -> str:
    """去除空行并将制表符替换为空格。"""
    lines = text.splitlines()
    cleaned_lines = [
        line.replace("\t", "    ") for line in lines if len(line.strip()) > 0
    ]
    return "\n".join(cleaned_lines)


def set_font(doc: Document) -> None:
    """设置全局字体为 Microsoft YaHei。"""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    rFonts = font.element.rPr.rFonts
    rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_title_page(doc: Document, title: str, version: str, company: str) -> None:
    # 计算剩余空间并将公司信息放置在接近底部的位置
    section = doc.sections[0]
    page_height = (
        section.page_height.cm
        - section.top_margin.cm
        - section.bottom_margin.cm
        - 0.4  # section.header
        - 0.4  # section.footer
    )
    current_height = (
        2.5  # 假设标题行高约为2.5cm
        + 0.5  # 假设版本行高约为0.5cm
        + 0.5  # 假设“源代码”行高约为0.5cm
        + 8  # 手动再调整一下
    )  # 假设段间距为1.5cm

    remaining_space = (page_height - current_height) / 2

    # 添加空白段落以填充剩余空间
    blank_paragraph = doc.add_paragraph("")
    blank_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    blank_paragraph.paragraph_format.space_after = Pt(
        remaining_space * 28.35
    )  # cm to pt conversion

    # 创建标题页
    title_paragraph = doc.add_paragraph(title)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(26)  # 一号字体大小约为26pt

    # 添加版本信息
    version_paragraph = doc.add_paragraph(version)
    version_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    version_run = version_paragraph.runs[0]
    version_run.font.size = Pt(14)

    # 添加“源代码”字样
    source_code_paragraph = doc.add_paragraph("源代码")
    source_code_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    source_code_run = source_code_paragraph.runs[0]
    source_code_run.font.size = Pt(14)

    # 添加空白段落以填充剩余空间
    blank_paragraph = doc.add_paragraph("")
    blank_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    blank_paragraph.paragraph_format.space_after = Pt(
        remaining_space * 28.35
    )  # cm to pt conversion

    # 添加公司信息
    company_paragraph = doc.add_paragraph(company)
    company_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    company_run = company_paragraph.runs[0]
    company_run.font.size = Pt(14)

    # 添加分页符
    doc.add_page_break()


def add_header_footer(doc: Document, title: str, version: str) -> None:
    # 设置页边距为最窄
    section = doc.sections[0]
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(0.6)
    section.right_margin = Cm(0.6)
    section.header_distance = Cm(0.1)
    section.footer_distance = Cm(0.1)

    # 设置页眉高度
    header = section.header
    header.height = Cm(0.2)
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_paragraph.paragraph_format.space_before = Pt(0)
    header_paragraph.paragraph_format.space_after = Pt(0)

    # 添加页眉内容
    header_run = header_paragraph.add_run(f"{title} {version}\t\t\t")
    header_run.font.size = Pt(10)

    # 添加当前页码到页眉右侧
    run = header_paragraph.add_run()
    fldSimple = OxmlElement("w:fldSimple")
    fldSimple.set(qn("w:instr"), " PAGE ")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fldSimple.append(r)
    run._r.append(fldSimple)

    # 添加斜杠到页眉右侧
    run = header_paragraph.add_run("/")
    run.font.size = Pt(10)

    # 添加总页码到页眉右侧
    run = header_paragraph.add_run()
    fldSimple = OxmlElement("w:fldSimple")
    fldSimple.set(qn("w:instr"), " NUMPAGES ")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fldSimple.append(r)
    run._r.append(fldSimple)

    # 设置页脚高度
    footer = section.footer
    footer.height = Cm(0.2)
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_paragraph.paragraph_format.space_before = Pt(0)
    footer_paragraph.paragraph_format.space_after = Pt(0)

    # 添加页脚内容
    footer_run = footer_paragraph.add_run(f"{title} {version}\t\t\t")
    footer_run.font.size = Pt(10)

    # 添加当前页码到页脚右侧
    run = footer_paragraph.add_run()
    fldSimple = OxmlElement("w:fldSimple")
    fldSimple.set(qn("w:instr"), " PAGE ")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fldSimple.append(r)
    run._r.append(fldSimple)

    # 添加斜杠到页脚右侧
    run = footer_paragraph.add_run("/")
    run.font.size = Pt(10)

    # 添加总页码到页脚右侧
    run = footer_paragraph.add_run()
    fldSimple = OxmlElement("w:fldSimple")
    fldSimple.set(qn("w:instr"), " NUMPAGES ")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fldSimple.append(r)
    run._r.append(fldSimple)

    # 添加制表符并设置其位置
    tab_stops = header_paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(20), WD_TAB_ALIGNMENT.RIGHT)
    tab_stops = footer_paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(20), WD_TAB_ALIGNMENT.RIGHT)


def add_main_content(
    doc: Document,
    files: list[Path],
    include_line_numbers: bool,
    max_lines: int = 2000,
) -> None:
    total_lines = 0  # 初始化总行数计数器

    for file_index, file_path in enumerate(files):
        if total_lines > max_lines:
            logger.info("已达到行数限制 %d", max_lines)
            break

        full_path = Path(file_path)
        try:
            with open(full_path, encoding="utf-8") as file:
                content = file.read()
                content = remove_consecutive_blank_lines(content)
                lines = content.splitlines()

                if include_line_numbers:
                    numbered_lines = [
                        f"{line_no}: {line}"
                        for line_no, line in enumerate(lines, start=total_lines)
                    ]
                    content = "\n".join(numbered_lines)

                doc.add_paragraph(content)
                total_lines += len(lines)  # 更新总行数计数器
                logger.debug("加入第 %d 个文件 %s", file_index + 1, full_path)

        except Exception as e:
            logger.error("无法读取文件 %s: %s", full_path, e)


def add_content_to_docx(
    doc: Document,
    title: str,
    version: str,
    company: str,
    files: list[Path],
    include_line_numbers: bool,
    max_lines: int,
) -> None:
    # 设置全局字体
    set_font(doc)

    # 添加页眉和页脚
    add_header_footer(doc, title, version)

    # 添加标题页
    add_title_page(doc, title, version, company)

    # 添加正文内容
    add_main_content(doc, files, include_line_numbers, max_lines=max_lines)


def _generate_builtin_doc(
    group: dict[str, Any],
    matched_files: list[Path],
    base_path: Path,
    output_file: Path,
) -> None:
    """使用内置 python-docx 在进程内生成 DOCX（纯 MIT/BSD，无 GPL 依赖）。"""
    title = group.get("title", "未命名文档")
    version = group.get("version", "1.0")
    company = group.get("company", "未知公司")
    include_line_numbers = group.get("lineNumber", False)
    max_lines = group.get("maxLines", 2000)

    # 将扫描得到的相对路径解析为基于 cwd 的绝对路径后读取
    resolved_files = [
        (base_path / f).resolve() if not Path(f).is_absolute() else Path(f).resolve()
        for f in matched_files
    ]

    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    add_content_to_docx(
        doc,
        title,
        version,
        company,
        resolved_files,
        include_line_numbers,
        max_lines=max_lines,
    )
    doc.save(output_file)
    logger.info("DOCX 文件已生成: %s", output_file)


def _generate_external_doc(
    matched_files: list[Path],
    base_path: Path,
    output_file: Path,
    group: dict[str, Any],
) -> None:
    """通过子进程委托 copyright-code-extractor (GPL-3.0) 生成 DOCX。"""
    if importlib.util.find_spec("copyright_code_extractor") is None:
        raise RuntimeError(
            "external 后端依赖 copyright-code-extractor 未安装。\n"
            '请通过 `uv tool install "pysample[external]"` 安装后再使用 --backend external。'
        )

    # 暂存命中文件到临时目录（保留目录结构），以便 copyright-code-extractor
    # 能通过目录扫描方式处理（该工具不支持传入文件列表）
    temp_dir = _prepare_temp_directory(matched_files, base_path)

    try:
        # 生成 copyright-code-extractor 配置文件
        config_file = _write_extractor_config(temp_dir, output_file, group)

        # 通过子进程委托给 copyright-code-extractor（GPL-3.0）。
        # 使用 sys.executable -m 确保在任何 uvx / uv tool / pip 等
        # 安装方式下都能找到正确的解释器环境中的包。
        cmd = [
            sys.executable,
            "-m",
            "copyright_code_extractor.cli",
            str(temp_dir),
            "-c",
            str(config_file),
        ]
        logger.info("正在调用 copyright-code-extractor...")
        logger.debug("命令: %s", " ".join(cmd))

        result = subprocess.run(cmd, capture_output=True, text=True)

        if result.returncode != 0:
            logger.error(
                "copyright-code-extractor 执行失败 (exit code %d):\n%s",
                result.returncode,
                result.stderr,
            )
            raise RuntimeError(f"copyright-code-extractor failed: {result.stderr}")

        if result.stdout:
            logger.info("%s", result.stdout.strip())

        logger.info("DOCX 文件已生成: %s", output_file)

    finally:
        # 清理临时目录
        shutil.rmtree(temp_dir, ignore_errors=True)
        logger.debug("已清理临时目录: %s", temp_dir)


def generate(verbose: bool, backend: str) -> None:
    """遍历配置中的 group，扫描源码并按所选后端生成 DOCX。"""
    if verbose:
        logging.getLogger().setLevel(logging.DEBUG)
        logger.info("详细日志已启用")
    else:
        logging.getLogger().setLevel(logging.INFO)

    config = read_config(CONFIG_PATH)

    for i, group in enumerate(config.get("group", [])):
        if verbose:
            logger.debug("处理第 %d 个组 %s", i + 1, group.get("title", "(未命名)"))

        base_dir = group.get("cwd", "")
        patterns = group.get("patterns", ["**"])
        extensions = group.get("extensions", [])
        ignore_dirs = group.get("ignore", [])
        title = group.get("title", "未命名文档")
        version = group.get("version", "1.0")

        matched_files = scan_files(
            base_dir,
            patterns,
            extensions,
            ignore_dirs,
            title=title,
            verbose=verbose,
        )
        logger.info("文件扫描完成，共 %d 个文件", len(matched_files))

        if not matched_files:
            logger.warning("%s: 未匹配到任何文件，跳过", title)
            continue

        output_dir = Path("./.bring-it/sample")
        output_dir.mkdir(parents=True, exist_ok=True)
        output_file = output_dir / f"{title}_{version}.docx"

        base_path = Path(base_dir).resolve()

        if backend == "builtin":
            _generate_builtin_doc(group, matched_files, base_path, output_file)
        elif backend == "external":
            if group.get("company") or group.get("lineNumber"):
                logger.warning(
                    "%s: external 后端不支持 company/lineNumber，相关配置已忽略", title
                )
            _generate_external_doc(matched_files, base_path, output_file, group)
        else:
            raise ValueError(f"未知后端: {backend}，可选 builtin / external")


def detect_languages(base_dir: str = ".") -> tuple[list[str], list[str]]:
    """扫描项目中常见源码格式，返回 (检测到的语言名列表, 去重扩展名列表)。

    判定为「已检测到」的条件：某语言的文件数 >= 2 或 总字节数 >= 1024 (1kB)。
    """
    base_path = Path(base_dir)
    if not base_path.exists() or not base_path.is_dir():
        return [], []

    # 扩展名 -> 语言名 反向映射
    ext_to_lang: dict[str, str] = {}
    for lang, exts in LANGUAGE_MAP.items():
        for ext in exts:
            ext_to_lang[ext] = lang

    lang_count: dict[str, int] = {}
    lang_size: dict[str, int] = {}

    for root, dirs, files in os.walk(base_path):
        # 原地修剪，跳过垃圾/产物目录以加快遍历
        dirs[:] = [d for d in dirs if d not in SKIP_DIRS]
        for fname in files:
            ext = fname.rsplit(".", 1)[-1].lower() if "." in fname else ""
            lang = ext_to_lang.get(ext)
            if lang is None:
                continue
            lang_count[lang] = lang_count.get(lang, 0) + 1
            with contextlib.suppress(OSError):
                lang_size[lang] = (
                    lang_size.get(lang, 0) + (base_path / root / fname).stat().st_size
                )

    detected_langs: list[str] = []
    detected_exts: list[str] = []
    for lang, exts in LANGUAGE_MAP.items():
        count = lang_count.get(lang, 0)
        size = lang_size.get(lang, 0)
        if count >= 2 or size >= 1024:
            detected_langs.append(lang)
            for ext in exts:
                if ext not in detected_exts:
                    detected_exts.append(ext)

    logger.info(
        "自动探测到 %d 种语言: %s",
        len(detected_langs),
        ", ".join(detected_langs) or "(无)",
    )
    return detected_langs, detected_exts


def run_init() -> None:
    """交互式收集信息并写出 .bring-it/sample.config.json。"""
    logger.info("开始初始化 pysample 配置...")

    _, detected_exts = detect_languages(".")

    title = typer.prompt("软件名称 (title)", default=Path.cwd().name)
    version = typer.prompt("软件版本 (version)", default="V1.0")

    if detected_exts:
        logger.info("自动探测到的源码扩展名: %s", ", ".join(detected_exts))
        use_detected = typer.confirm("是否使用上述自动探测的扩展名？", default=True)
        extensions = list(detected_exts) if use_detected else []
    else:
        logger.warning("未自动探测到任何源码，请手动补充扩展名。")
        extensions = []

    extra = typer.prompt("额外的文件后缀（逗号分隔，可留空）", default="").strip()
    for e in extra.split(","):
        e = e.strip().lstrip(".").lower()
        if e and e not in extensions:
            extensions.append(e)

    if not extensions:
        logger.warning("未指定任何扩展名，初始化后将无法扫描到文件。")

    cwd = typer.prompt("基准目录 (cwd)", default=".")

    ignore_input = typer.prompt(
        "忽略目录（逗号分隔，可留空，.gitignore 已自动生效）",
        default="",
    ).strip()
    ignore = [d.strip() for d in ignore_input.split(",") if d.strip()]

    config: dict[str, Any] = {
        "group": [
            {
                "cwd": cwd,
                "patterns": ["**"],
                "extensions": extensions,
                "ignore": ignore,
                "title": title,
                "version": version,
            }
        ]
    }

    CONFIG_PATH.parent.mkdir(parents=True, exist_ok=True)
    with open(CONFIG_PATH, "w", encoding="utf-8") as f:
        json.dump(config, f, indent=2, ensure_ascii=False)

    logger.info("已生成配置文件: %s", CONFIG_PATH.resolve())


@app.callback(invoke_without_command=True)
def main(
    ctx: typer.Context,
    verbose: bool = typer.Option(False, "--verbose", help="启用详细日志输出"),
    backend: str = typer.Option(
        "builtin",
        "--backend",
        help="文档生成后端: builtin(内置 python-docx, 默认, 纯 MIT) 或 external(委托 copyright-code-extractor, GPL-3.0)",
    ),
) -> None:
    """生成软件著作权源代码文档（bring-it/sample 兼容）。

    若未发现 .bring-it/sample.config.json 会交互式提示初始化。
    """
    # 有子命令（如 init）时交给子命令自行处理，避免重复执行文档生成
    if ctx.invoked_subcommand is not None:
        return

    if backend not in ("builtin", "external"):
        logger.error("未知后端: %s，可选 builtin / external", backend)
        raise typer.Exit(code=1)

    if not CONFIG_PATH.exists():
        if typer.confirm(
            "未发现 .bring-it/sample.config.json，是否现在初始化？(Y/n)",
            default=True,
        ):
            run_init()
        else:
            logger.error(
                "未初始化配置文件，无法继续生成文档。可运行 `pysample init` 进行初始化。"
            )
            raise typer.Exit(code=1)

    generate(verbose, backend)


@app.command()
def init() -> None:
    """交互式初始化/重新初始化 .bring-it/sample.config.json。"""
    if CONFIG_PATH.exists() and not typer.confirm(
        "已存在 .bring-it/sample.config.json，是否重新初始化？(y/N)",
        default=False,
    ):
        logger.info("已取消重新初始化，未做任何改动。")
        raise typer.Exit()
    run_init()


if __name__ == "__main__":
    app()
